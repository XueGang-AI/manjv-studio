#!/usr/bin/env python3
"""解析 DOCX 文件 → 输出结构化 JSON"""
import json, os
from docx import Document

FILES = [
    ("AI漫剧创作专业版提示词.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/AI漫剧创作/AI漫剧创作专业版提示词.docx",
     "AI漫剧创作"),
    ("12组电影级组合运镜提示词.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/AI运镜提示词/12组电影级组合运镜提示词.docx",
     "AI运镜提示词"),
    ("十大经典运镜教程.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/AI运镜提示词/十大经典运镜教程.docx",
     "AI运镜提示词"),
    ("特效运镜.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/AI运镜提示词/特效运镜.docx",
     "AI运镜提示词"),
    ("组合运镜.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/AI运镜提示词/组合运镜.docx",
     "AI运镜提示词"),
    ("组合运镜2.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/AI运镜提示词/组合运镜2.docx",
     "AI运镜提示词"),
    ("AI漫剧16个运镜提示词模板.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/分镜提示词＋AI指令整理/小说漫剧提示词/AI漫剧16个运镜提示词模板.docx",
     "小说漫剧提示词"),
    ("提示词大全（在夸克）.docx",
     "/Users/xuegang/Desktop/My Project/AI漫剧制作/提示词大全（在夸克）.docx",
     "AI漫剧制作"),
]

OUTPUT = "/Users/xuegang/Desktop/My Project/manjv-studio/scripts/output/docx_parsed.json"
results = []

for name, path, source_dir in FILES:
    entry = {"file": name, "source_dir": source_dir}
    if not os.path.exists(path):
        entry["status"] = "NOT_FOUND"
        results.append(entry)
        print(f"❌ {name}: NOT FOUND")
        continue
    
    try:
        doc = Document(path)
        paragraphs = []
        tables = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "style": para.style.name if para.style else "Normal",
                    "text": text
                })
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        entry["status"] = "OK"
        entry["paragraph_count"] = len(paragraphs)
        entry["table_count"] = len(tables)
        entry["paragraphs"] = paragraphs
        entry["tables"] = tables
        results.append(entry)
        
        para_chars = sum(len(p["text"]) for p in paragraphs)
        print(f"✅ {name}: {len(paragraphs)} paragraphs, {len(tables)} tables, {para_chars} chars")
        
    except Exception as e:
        entry["status"] = "ERROR"
        entry["error"] = str(e)
        results.append(entry)
        print(f"❌ {name}: ERROR - {e}")

with open(OUTPUT, 'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)

print(f"\n📄 Output: {OUTPUT}")
